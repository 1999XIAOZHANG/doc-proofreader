import os
import uuid
from docx import Document
from typing import Dict, List, Optional
from datetime import datetime


class DocumentStore:
    """文档存储和管理"""
    
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = upload_dir
        self.documents: Dict[str, dict] = {}  # 内存存储文档信息
        os.makedirs(upload_dir, exist_ok=True)
    
    def save_document(self, document_id: str, filename: str, file_content: bytes) -> str:
        """保存上传的文档，返回document_id"""
        file_path = os.path.join(self.upload_dir, f"{document_id}.docx")
        
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        self.documents[document_id] = {
            "filename": filename,
            "file_path": file_path,
            "created_at": datetime.now(),
            "paragraphs": self._extract_paragraphs(file_path),
            "corrections": []
        }
        return document_id
    
    def _extract_paragraphs(self, file_path: str) -> List[str]:
        """从Word文档提取段落文本"""
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        return paragraphs
    
    def get_paragraphs(self, document_id: str) -> Optional[List[str]]:
        """获取文档的所有段落"""
        doc_info = self.documents.get(document_id)
        if doc_info:
            return doc_info["paragraphs"]
        return None
    
    def get_file_path(self, document_id: str) -> Optional[str]:
        """获取文档文件路径"""
        doc_info = self.documents.get(document_id)
        if doc_info:
            return doc_info["file_path"]
        return None
    
    def save_paragraphs(self, document_id: str, paragraphs: List[str]):
        """保存段落列表"""
        if document_id in self.documents:
            self.documents[document_id]["paragraphs"] = paragraphs

    def save_corrections(self, document_id: str, corrections: list):
        """保存校对结果"""
        if document_id in self.documents:
            self.documents[document_id]["corrections"] = corrections

    def get_corrections(self, document_id: str) -> list:
        """获取校对结果"""
        doc_info = self.documents.get(document_id)
        if doc_info:
            return doc_info.get("corrections", [])
        return []

    def update_paragraph(self, document_id: str, paragraph_index: int, new_text: str):
        """更新段落内容"""
        if document_id in self.documents:
            self.documents[document_id]["paragraphs"][paragraph_index] = new_text
    
    def save_modified_document(self, document_id: str, output_path: str):
        """保存修改后的文档"""
        original_path = self.get_file_path(document_id)
        if not original_path:
            raise ValueError("Document not found")
        
        doc = Document(original_path)
        paragraphs = self.get_paragraphs(document_id)
        
        # 更新段落
        para_idx = 0
        for para in doc.paragraphs:
            if para.text.strip() and para_idx < len(paragraphs):
                # 保留原有格式，只修改文本
                for run in para.runs:
                    run.text = ""  # 清空原有runs
                if len(para.runs) > 0:
                    para.runs[0].text = paragraphs[para_idx]
                else:
                    para.add_run(paragraphs[para_idx])
                para_idx += 1
        
        doc.save(output_path)


# 全局文档存储实例
doc_store = DocumentStore()
